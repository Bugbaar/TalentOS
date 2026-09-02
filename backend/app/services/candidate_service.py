"""Candidate persistence and resume ingestion services."""

import logging
import re
import uuid
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader
from sqlalchemy import or_, select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.core.exceptions import NotFoundError, ValidationError
from app.models.candidate import Candidate, Resume
from app.models.job import JobOpening
from app.schemas.ai import OutreachEmailResponse
from app.schemas.candidate import CandidateCreate
from app.services.ai_engine.factory import get_ai_provider

logger = logging.getLogger(__name__)

# Maximum file size for resume uploads (15MB)
MAX_FILE_SIZE = 15 * 1024 * 1024

# Allowed file extensions and MIME types
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
ALLOWED_MIME_TYPES = {
    ".pdf": {"application/pdf"},
    ".docx": {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
    ".txt": {"text/plain", "text/markdown"},
    ".md": {"text/plain", "text/markdown"},
}


async def get_candidates(
    db: AsyncSession,
    skip: int = 0,
    limit: int = 100,
    skill_filter: str | None = None,
    search: str | None = None,
) -> list[Candidate]:
    """Return candidates filtered by skill and free-text search."""

    query = (
        select(Candidate)
        .order_by(Candidate.created_at.desc())
        .offset(max(skip, 0))
        .limit(min(limit, 1000))
    )
    if skill_filter:
        query = query.where(Candidate.skills.contains([skill_filter]))
    if search:
        term = f"%{search.strip()}%"
        query = query.where(
            or_(
                Candidate.first_name.ilike(term),
                Candidate.last_name.ilike(term),
                Candidate.email.ilike(term),
            )
        )
    result = await db.execute(query)
    return list(result.scalars().all())


async def get_candidate_by_id(
    db: AsyncSession, candidate_id: uuid.UUID
) -> Candidate | None:
    """Return one candidate with related resumes and applications."""

    result = await db.execute(
        select(Candidate)
        .options(selectinload(Candidate.resumes), selectinload(Candidate.applications))
        .where(Candidate.id == candidate_id)
    )
    return result.scalar_one_or_none()


async def get_candidate_by_id_or_raise(
    db: AsyncSession, candidate_id: uuid.UUID
) -> Candidate:
    """Return a candidate or raise a NotFoundError."""

    candidate = await get_candidate_by_id(db, candidate_id)
    if candidate is None:
        raise NotFoundError("Candidate", str(candidate_id))
    return candidate


async def create_candidate(db: AsyncSession, candidate_in: CandidateCreate) -> Candidate:
    """Persist and return a candidate created from validated input."""

    candidate = Candidate(**candidate_in.model_dump())
    db.add(candidate)
    await db.commit()
    await db.refresh(candidate)
    return candidate


async def update_candidate(
    db: AsyncSession,
    candidate_id: uuid.UUID,
    update_data: dict[str, Any],
) -> Candidate:
    """Update an existing candidate's fields."""

    candidate = await get_candidate_by_id_or_raise(db, candidate_id)
    for key, value in update_data.items():
        if value is not None:
            setattr(candidate, key, value)
    await db.commit()
    await db.refresh(candidate)
    return candidate


async def delete_candidate(db: AsyncSession, candidate_id: uuid.UUID) -> None:
    """Delete a candidate and all related records."""

    candidate = await get_candidate_by_id_or_raise(db, candidate_id)
    await db.delete(candidate)
    await db.commit()


async def parse_and_create_candidate_from_resume(
    db: AsyncSession, file_bytes: bytes, filename: str
) -> Candidate:
    """Extract resume text, parse it with the configured AI provider, and persist both records."""

    if not file_bytes:
        raise ValidationError("The uploaded file is empty")
    if len(file_bytes) > MAX_FILE_SIZE:
        raise ValidationError(
            f"File exceeds maximum size of {MAX_FILE_SIZE // (1024 * 1024)}MB"
        )

    text = extract_text_from_file(file_bytes, filename)
    parsed = await get_ai_provider().parse_resume(text)
    first_name, last_name = parsed.first_name.strip(), parsed.last_name.strip()
    if not first_name or not last_name:
        first_name, last_name = _name_from_filename(filename)
    candidate = Candidate(
        first_name=first_name or "Unknown",
        last_name=last_name or "Candidate",
        email=parsed.email or f"resume-{uuid.uuid4().hex}@invalid.local",
        phone=parsed.phone,
        headline=parsed.headline,
        experience_years=parsed.experience_years,
        skills=parsed.skills,
    )
    resume = Resume(
        candidate=candidate,
        raw_text=text,
        file_url=filename,
        parsed_skills=parsed.skills,
        work_history=parsed.work_history,
        education=parsed.education,
    )
    db.add(candidate)
    await db.commit()
    persisted = await get_candidate_by_id(db, candidate.id)
    if persisted is None:
        raise RuntimeError("Candidate could not be loaded after resume ingestion")
    logger.info("Ingested resume for candidate %s", candidate.id)
    return persisted


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Extract readable text from PDF, DOCX, or plain-text resume bytes.

    The format is selected from the file extension rather than trusting the
    client-provided content type. Unsupported formats and malformed documents
    raise ``ValueError`` so API callers can return a useful 4xx response.
    """

    if not file_bytes:
        raise ValidationError("The uploaded file is empty")
    extension = Path(filename or "").suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise ValidationError(
            f"Unsupported file type '{extension}'; upload a PDF, DOCX, TXT, or Markdown file"
        )
    if extension == ".pdf":
        try:
            reader = PdfReader(BytesIO(file_bytes))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as exc:
            raise ValidationError("Unable to read the PDF document") from exc
    elif extension == ".docx":
        try:
            document = Document(BytesIO(file_bytes))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            tables = [
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            ]
            text = "\n".join(part for part in paragraphs + tables if part.strip())
        except Exception as exc:
            raise ValidationError("Unable to read the DOCX document") from exc
    elif extension in {".txt", ".md"}:
        try:
            text = file_bytes.decode("utf-8-sig", errors="strict")
        except UnicodeDecodeError:
            text = file_bytes.decode("latin-1", errors="replace")
    else:
        raise ValidationError("Unsupported file type; upload a PDF, DOCX, TXT, or Markdown file")
    if not text.strip():
        raise ValidationError("The uploaded document contains no extractable text")
    return text.strip()


def _name_from_filename(filename: str) -> tuple[str, str]:
    stem = Path(filename or "resume").stem
    words = re.findall(r"[A-Za-z][A-Za-z'-]*", stem)
    return (words[0], words[1]) if len(words) >= 2 else ("", "")


async def generate_outreach_email(
    db: AsyncSession,
    candidate_id: uuid.UUID,
    job_id: uuid.UUID | None,
    tone: str,
    company: str,
) -> OutreachEmailResponse:
    """Generate personalized outreach for a candidate and optional job."""

    candidate = await get_candidate_by_id_or_raise(db, candidate_id)
    job = await db.get(JobOpening, job_id) if job_id else None
    if job_id and job is None:
        raise NotFoundError("Job", str(job_id))
    candidate_data = {
        "first_name": candidate.first_name,
        "last_name": candidate.last_name,
        "headline": candidate.headline,
        "experience_years": candidate.experience_years,
        "skills": candidate.skills or [],
    }
    job_data = (
        {
            "title": job.title,
            "description": job.description,
            "required_skills": job.required_skills or [],
        }
        if job
        else None
    )
    return await get_ai_provider().generate_outreach_email(
        candidate_data, job_data, tone, company
    )
